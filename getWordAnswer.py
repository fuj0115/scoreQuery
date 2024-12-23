from docx import Document
import os
def find_doc_path(file_name, root_folder='D:/资料'):
    """
    递归遍历指定文件夹下的所有 .docx 文件，找到文件名模糊匹配 file_name 的第一个文件，返回其路径。
    :param file_name: 要查找的文件名（可以是部分文件名）
    :param root_folder: 起始搜索的文件夹路径，默认为 'D:/资料文件夹'
    :return: 找到的文件路径，如果没有找到则返回 None
    """
    for foldername, subfolders, filenames in os.walk(root_folder):
        for filename in filenames:
            if filename.lower().endswith('.docx') and file_name.lower() in filename.lower():
                return os.path.join(foldername, filename)
    return None


def search_and_extract_answer(doc_path, custom_string):
    doc = Document(doc_path)
    found = False
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if custom_string in text:
            found = True
            print(f"找到题目: {text}")  # 调试信息
            for next_paragraph in doc.paragraphs[i + 1:]:
                next_text = next_paragraph.text.strip()
                print(f"检查答案行: {next_text}")  # 调试信息
                if  "正确答案" in next_text :
                    print('找到答案行')
                    print(next_text.split("正确答案是：")[1].strip())
                    return next_text.split("正确答案是：")[1].strip()
    if not found:
        return "未找到相关题目"

# 使用示例
doc_path = find_doc_path('10775-经济法')
if doc_path:
    print(f"找到文件路径: {doc_path}")
else:
    print("未找到文件")

# 打印文档内容以调试
doc = Document(doc_path)
for i, paragraph in enumerate(doc.paragraphs):
     if "根据《政府采购法实施条例》，关于列入集中采购目录的项目，说法不正确的有" in paragraph.text.strip():
        print(f"找到题目: {paragraph.text.strip()}")
    #print(f"段落 {i+1}: {paragraph.text.strip()}")

# 查找特定题目的答案
custom_string = "根据《政府采购法实施条例》，关于列入集中采购目录的项目，说法不正确的有"
answer = search_and_extract_answer(doc_path, custom_string)
if answer:
    print(f"题目: {custom_string}\n答案: {answer}")
else:
    print(f"未找到题目: {custom_string}")